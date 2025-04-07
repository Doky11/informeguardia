import streamlit as st
from docx import Document
from io import BytesIO
from datetime import datetime

def nota_a_letra(nota):
    if nota >= 9:
        return 'A'
    elif nota >= 7:
        return 'B'
    elif nota >= 5:
        return 'C'
    else:
        return 'D'

st.set_page_config(page_title="Informe Guardia Naval", layout="wide")

st.title("Informe Personal de la Guardia")

st.header("Encabezado")
informante = st.text_input("Informante")
fecha = st.date_input("Fecha", format="DD.MM.YYYY")
alumno = st.text_input("Alumno")
puesto = st.text_input("Puesto")

st.header("Evaluación por Categoría")
categorias = [
    "POLICÍA", "DISCIPLINA", "INTERES", "RESPONSABILIDAD", "INICIATIVA",
    "CONFIANZA EN SI MISMO", "ACTITUD CON LOS SUBORDINADOS", "ACTITUD CON EL MANDO",
    "COMPETENCIA / EFICACIA", "TRATO", "RESISTENCIA A LA FATIGA"
]

resultados = {}
total_nota = 0

for categoria in categorias:
    st.subheader(categoria)
    checks = [st.checkbox(f"{categoria} - Pregunta {i}", key=f"{categoria}_{i}") for i in range(1, 7)]
    num_check = sum(checks)
    nota = round((num_check / 6) * 10, 2)
    letra = nota_a_letra(nota)
    observacion = st.text_area(f"Observaciones para {categoria}", value=f"{nota}")
    st.write(f"Nota numérica: **{nota}**, Nota en letra: **{letra}**")
    resultados[categoria] = (nota, letra, observacion)
    total_nota += nota

nota_media = round(total_nota / len(categorias), 2)
letra_media = nota_a_letra(nota_media)

st.subheader("Nota media")
st.write(f"Nota media numérica: **{nota_media}**")
st.write(f"Nota media en letra: **{letra_media}**")

observaciones_generales = st.text_area("Observaciones generales / Justificación")

if st.button("Generar informe"):
    doc = Document("INFORME EN BLANCO.docx")

    for tabla in doc.tables:
        for fila in tabla.rows:
            for i, celda in enumerate(fila.cells):
                texto = celda.text.upper()
                if "INFORMANTE" in texto and i < len(fila.cells) - 1:
                    fila.cells[i+1].text = informante
                elif "FECHA" in texto and i < len(fila.cells) - 1:
                    fila.cells[i+1].text = fecha.strftime('%d.%m.%Y')
                elif "ALUMNO" in texto and i < len(fila.cells) - 1:
                    fila.cells[i+1].text = alumno
                elif "PUESTO" in texto and i < len(fila.cells) - 1:
                    fila.cells[i+1].text = puesto
                elif "NOTA MEDIA" in texto:
                    fila.cells[i].text = f"Nota media: {nota_media} ({letra_media})"
                elif "OBSERVACIONES GENERAL" in texto and i < len(fila.cells) - 1:
                    fila.cells[i+1].text = observaciones_generales if i+1 < len(fila.cells) else ""

    try:
        if len(doc.tables) > 1:
            tabla = doc.tables[1]
            for i, categoria in enumerate(categorias):
                nota, letra, observacion_categoria = resultados[categoria]
                if i+1 < len(tabla.rows):
                    fila = tabla.rows[i+1].cells
                    if len(fila) >= 6:
                        fila[5].text = observacion_categoria
                        if letra == 'A': fila[1].text = 'X'
                        elif letra == 'B': fila[2].text = 'X'
                        elif letra == 'C': fila[3].text = 'X'
                        else: fila[4].text = 'X' 
            nota, letra, observacion_categoria = resultados[categoria]
            fila = tabla.rows[i+1].cells if i+1 < len(tabla.rows) else []
            if len(fila) >= 6:
                fila[5].text = observacion_categoria
                if letra == 'A':
                    fila[1].text = 'X'
                elif letra == 'B':
                    fila[2].text = 'X'
                elif letra == 'C':
                    fila[3].text = 'X'
                else:
                    fila[4].text = 'X'
    except Exception as e:
        st.error(f"⚠️ Error al escribir en la tabla del documento: {e}")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    st.download_button(
        label="📄 Descargar informe Word",
        data=buffer,
        file_name=f"Informe_{alumno}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
