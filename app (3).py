import streamlit as st
from docx import Document
from io import BytesIO
from datetime import datetime

def nota_a_letra(nota):
    if nota >= 9:
        return 'A'
    elif nota >= 7:
        return 'B'
    elif nota >= 5:
        return 'C'
    else:
        return 'D'

st.set_page_config(page_title="Informe Guardia Naval", layout="wide")

st.title("Informe Personal de la Guardia")

st.header("Encabezado")
informante = st.text_input("Informante")
fecha = st.date_input("Fecha", format="DD.MM.YYYY")
alumno = st.text_input("Alumno")
puesto = st.text_input("Puesto")

st.header("Evaluación por Categoría")
categorias = [
    "POLICÍA", "DISCIPLINA", "INTERES", "RESPONSABILIDAD", "INICIATIVA",
    "CONFIANZA EN SI MISMO", "ACTITUD CON LOS SUBORDINADOS", "ACTITUD CON EL MANDO",
    "COMPETENCIA / EFICACIA", "TRATO", "RESISTENCIA A LA FATIGA"
]

resultados = {}
total_nota = 0

for categoria in categorias:
    st.subheader(categoria)
    checks = [st.checkbox(f"{categoria} - Pregunta {i}", key=f"{categoria}_{i}") for i in range(1, 7)]
    num_check = sum(checks)
    nota = round((num_check / 6) * 10, 2)
    letra = nota_a_letra(nota)
    observacion = st.text_area(f"Observaciones para {categoria}", value=f"{nota}")
    st.write(f"Nota numérica: **{nota}**, Nota en letra: **{letra}**")
    resultados[categoria] = (nota, letra, observacion)
    total_nota += nota

nota_media = round(total_nota / len(categorias), 2)
letra_media = nota_a_letra(nota_media)

st.subheader("Nota media")
st.write(f"Nota media numérica: **{nota_media}**")
st.write(f"Nota media en letra: **{letra_media}**")

observaciones_generales = st.text_area("Observaciones generales / Justificación")

if st.button("Generar informe"):
    doc = Document("INFORME EN BLANCO.docx")

    # Rellenar encabezado
    for tabla in doc.tables:
        for fila in tabla.rows:
            for i, celda in enumerate(fila.cells):
                texto = celda.text.upper().strip()
                if texto.startswith("INFORMANTE") and i + 1 < len(fila.cells):
                    fila.cells[i+1].text = informante
                elif texto.startswith("FECHA") and i + 1 < len(fila.cells):
                    fila.cells[i+1].text = fecha.strftime('%d.%m.%Y')
                elif texto.startswith("ALUMNO") and i + 1 < len(fila.cells):
                    fila.cells[i+1].text = alumno
                elif texto.startswith("PUESTO") and i + 1 < len(fila.cells):
                    fila.cells[i+1].text = puesto

    # Buscar tabla con las categorías y rellenar
    for tabla in doc.tables:
        for i, fila in enumerate(tabla.rows):
            celdas = fila.cells
            if len(celdas) >= 6 and celdas[0].text.strip().upper() in categorias:
                categoria = celdas[0].text.strip().upper()
                if categoria in resultados:
                    nota, letra, observacion = resultados[categoria]
                    celdas[5].text = observacion  # observaciones
                    if letra == 'A':
                        celdas[1].text = 'X'
                    elif letra == 'B':
                        celdas[2].text = 'X'
                    elif letra == 'C':
                        celdas[3].text = 'X'
                    else:
                        celdas[4].text = 'X'

    # Rellenar nota media y observaciones generales en tabla final
    for tabla in doc.tables:
        for fila in tabla.rows:
            for i, celda in enumerate(fila.cells):
                texto = celda.text.upper()
                if "NOTA MEDIA" in texto:
                    celda.text = f"Nota media: {nota_media} ({letra_media})"
                elif "OBSERVACIONES GENERAL" in texto and i + 1 < len(fila.cells):
                    fila.cells[i+1].text = observaciones_generales

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    st.download_button(
        label="📄 Descargar informe Word",
        data=buffer,
        file_name=f"Informe_{alumno}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
