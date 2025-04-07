import streamlit as st
from docx import Document
from io import BytesIO
from datetime import datetime

# Diccionario para convertir notas numéricas en letras
def nota_a_letra(nota):
    if nota >= 9:
        return 'A'
    elif nota >= 7:
        return 'B'
    elif nota >= 5:
        return 'C'
    else:
        return 'D'

st.set_page_config(page_title="Informe Guardia Naval", layout="wide")
st.markdown("""
    <style>
    body {
        background-color: #e0f0ff;
    }
    .main {
        background-color: #f2f9ff;
        border: 2px solid #00557f;
        border-radius: 10px;
        padding: 20px;
    }
    h1, h2, h3 {
        color: #00334e;
    }
    </style>
""", unsafe_allow_html=True)

st.title("Informe Personal de la Guardia")

st.header("Encabezado")
informante = st.text_input("Informante")
fecha = st.date_input("Fecha", format="DD.MM.YYYY")
alumno = st.text_input("Alumno")
puesto = st.text_input("Puesto")

st.header("Evaluación por Categoría")
categorias = [
    "POLICÍA", "DISCIPLINA", "INTERES", "RESPONSABILIDAD", "INICIATIVA",
    "CONFIANZA EN SI MISMO", "ACTITUD CON LOS SUBORDINADOS", "ACTITUD CON EL MANDO",
    "COMPETENCIA / EFICACIA", "TRATO", "RESISTENCIA A LA FATIGA"
]

resultados = {}
total_nota = 0

for categoria in categorias:
    st.subheader(categoria)
    checks = []
    for i in range(1, 7):
        checks.append(st.checkbox(f"{categoria} - Pregunta {i}", key=f"{categoria}_{i}"))
    num_check = sum(checks)
    nota = round((num_check / 6) * 10, 2)
    letra = nota_a_letra(nota)
    st.write(f"Nota numérica: **{nota}**, Nota en letra: **{letra}**")
    resultados[categoria] = (nota, letra)
    total_nota += nota

nota_media = round(total_nota / len(categorias), 2)
letra_media = nota_a_letra(nota_media)

st.subheader("Nota media")
st.write(f"Nota media numérica: **{nota_media}**")
st.write(f"Nota media en letra: **{letra_media}**")

observaciones_generales = st.text_area("Observaciones generales / Justificación")

# Botón para generar el documento Word
if st.button("Generar informe"):
    doc = Document("INFORME EN BLANCO.docx")

    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                texto = celda.text
                if "INFORMANTE" in texto:
                    celda.text = f"INFORMANTE\n{informante}"
                elif "FECHA" in texto:
                    celda.text = f"FECHA\n{fecha.strftime('%d.%m.%Y')}"
                elif "ALUMNO" in texto:
                    celda.text = f"ALUMNO\n{alumno}"
                elif "PUESTO" in texto:
                    celda.text = f"PUESTO\n{puesto}"
                elif "Nota media:" in texto:
                    celda.text = f"Nota media: {nota_media} ({letra_media})"
                elif "OBSERVACIONES GENERAL" in texto:
                    celda.text = f"OBSERVACIONES GENERAL / JUSTIFICACIÓN\n{observaciones_generales}"

    # Asegurarse de que la tabla tenga suficientes filas y columnas
    try:
        tabla = doc.tables[1]
        for i, categoria in enumerate(categorias):
            nota, letra = resultados[categoria]
            fila = tabla.rows[i+1].cells
            if len(fila) >= 6:
                fila[5].text = str(nota)
                if letra == 'A': fila[1].text = 'X'
                elif letra == 'B': fila[2].text = 'X'
                elif letra == 'C': fila[3].text = 'X'
                else: fila[4].text = 'X'
    except IndexError:
        st.error("⚠️ Error al escribir en la tabla del documento. Asegúrate de que el archivo Word tiene el formato correcto y filas suficientes.")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    st.download_button(
        label="📄 Descargar informe Word",
        data=buffer,
        file_name=f"Informe_{alumno}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
